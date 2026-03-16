import os
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


REFERENCE_DOC = "MSWM_J_J_2026_28113.docx"
ASSET_DIR = Path(__file__).resolve().parent / ".msw_assets"
LEFT_IMAGE_NAME = "image2.jpeg"
RIGHT_IMAGE_NAME = "image1.png"

PAPER_YEAR = "2025"
ISSUE = "2"
START_PAGE = 2966


@dataclass(frozen=True)
class FormatConfig:
    template: str = "msw"  # "msw" | "ijrss" | "ijmie"
    volume: str = "36"
    paper_year: str = PAPER_YEAR
    paper_month: str = "March"
    issue: str = ISSUE
    start_page: int = START_PAGE

RED = RGBColor(192, 0, 0)
BLACK = RGBColor(0, 0, 0)
IJRSS_BLUE = RGBColor(0, 0, 255)

IJRSS_ISSN = "2249-2496"
IJRSS_IMPACT_FACTOR = "7.081"
IJRSS_HOMEPAGE = "http://www.ijmra.us"
IJRSS_EMAIL = "editorijmie@gmail.com"

IJMIE_ISSN = "2249-0558"
IJMIE_IMPACT_FACTOR = "7.119"

IJRSS_ORANGE_HEX = "ED7D31"
IJRSS_GRAY_HEX = "7F7F7F"

PAGE_OVERRIDES = {
    "Paper 2.docx": 5,
    "Paper 3.docx": 5,
}


def _docprops_page_count(filepath):
    try:
        with zipfile.ZipFile(filepath) as archive:
            app_xml = archive.read("docProps/app.xml").decode("utf-8", "ignore")
    except Exception:
        return 1

    match = re.search(r"<Pages>(\d+)</Pages>", app_xml)
    return int(match.group(1)) if match else 1


def get_page_count(filepath):
    name = Path(filepath).name
    if name in PAGE_OVERRIDES:
        return PAGE_OVERRIDES[name]
    return _docprops_page_count(filepath)


def ensure_reference_images(base_dir, reference_dir=None):
    reference_dir = reference_dir or base_dir
    reference_path = Path(reference_dir) / REFERENCE_DOC
    left_path = ASSET_DIR / LEFT_IMAGE_NAME
    right_path = ASSET_DIR / RIGHT_IMAGE_NAME

    if left_path.exists() and right_path.exists():
        return str(left_path), str(right_path)

    if not reference_path.exists():
        return "", ""

    ASSET_DIR.mkdir(exist_ok=True)
    with zipfile.ZipFile(reference_path) as archive:
        left_path.write_bytes(archive.read(f"word/media/{LEFT_IMAGE_NAME}"))
        right_path.write_bytes(archive.read(f"word/media/{RIGHT_IMAGE_NAME}"))
    return str(left_path), str(right_path)


def set_starting_page_number(doc, start_num):
    for section in doc.sections:
        pg_num_type = section._sectPr.find(qn("w:pgNumType"))
        if pg_num_type is None:
            pg_num_type = OxmlElement("w:pgNumType")
            section._sectPr.append(pg_num_type)
        pg_num_type.set(qn("w:start"), str(start_num))


def clear_story(story):
    for child in list(story._element):
        story._element.remove(child)


def set_cell_border(cell, edge, size, *, color="C00000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    border = tc_borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        tc_borders.append(border)

    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:color"), str(color))


def set_cell_margins(cell, top, start, bottom, end):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        margin = tc_mar.find(qn(f"w:{side}"))
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            tc_mar.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_table_full_width(table, section):
    # Expand the header table to the full page width by offsetting the left margin.
    tbl = table._tbl
    tbl_pr = getattr(tbl, "tblPr", None)
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        # tblPr should come before tblGrid/tr rows.
        tbl.insert(0, tbl_pr)

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(section.page_width)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(-int(section.left_margin)))
    tbl_ind.set(qn("w:type"), "dxa")


def ensure_header_clearance(section, minimum_top_margin=Inches(1.15)):
    if int(section.top_margin) < int(minimum_top_margin):
        section.top_margin = minimum_top_margin


def _iter_header_variants(section):
    for name in ("header", "first_page_header", "even_page_header"):
        part = getattr(section, name, None)
        if part is not None:
            yield part


def _iter_footer_variants(section):
    for name in ("footer", "first_page_footer", "even_page_footer"):
        part = getattr(section, name, None)
        if part is not None:
            yield part


def style_run(run, color=BLACK, bold=True):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.color.rgb = color


def style_run_custom(
    run,
    *,
    font_name="Times New Roman",
    color=BLACK,
    size=Pt(12),
    bold=False,
    italic=False,
    underline=False,
):
    run.bold = bool(bold)
    run.italic = bool(italic)
    run.underline = bool(underline)
    run.font.name = font_name
    run.font.size = size
    run.font.color.rgb = color


def style_paragraph(paragraph, alignment):
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0


def add_page_field(run):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    run._r.append(instruction)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run._r.append(separate)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(end)


def clean_blank_paragraphs(doc):
    # Remove truly empty paragraphs but keep anything that contains a drawing.
    for p in list(doc.paragraphs):
        has_drawing = any(
            any(el.tag.endswith("}drawing") for el in r._r.iter())
            for r in p.runs
        )
        if not has_drawing and not p.text.strip():
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)


def _populate_header_story(
    header,
    section,
    start_page,
    end_page,
    left_image,
    right_image,
    *,
    volume="36",
    paper_year=PAPER_YEAR,
    issue=ISSUE,
):
    clear_story(header)

    usable_width = section.page_width - section.left_margin - section.right_margin
    table = header.add_table(rows=1, cols=3, width=usable_width)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell = table.cell(0, 0)
    center_cell = table.cell(0, 1)
    right_cell = table.cell(0, 2)

    # Force column widths (Word often ignores cell.width unless columns are set).
    left_w = Inches(1.00)
    right_w = Inches(1.05)
    if usable_width < (left_w + right_w + Inches(2.6)):
        # Scale down side columns if the page is narrow / margins large.
        scale = int(usable_width) / int(left_w + right_w + Inches(2.6))
        scale = max(0.6, min(1.0, scale))
        left_w = int(left_w * scale)
        right_w = int(right_w * scale)
    center_w = usable_width - left_w - right_w

    table.columns[0].width = left_w
    table.columns[1].width = center_w
    table.columns[2].width = right_w
    left_cell.width = left_w
    center_cell.width = center_w
    right_cell.width = right_w

    for cell in (left_cell, center_cell, right_cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        set_cell_border(cell, "bottom", 18)

    # Swap: right image on left, left image on right (per user request).
    p_left = left_cell.paragraphs[0]
    style_paragraph(p_left, WD_ALIGN_PARAGRAPH.LEFT)
    if right_image and os.path.exists(right_image):
        p_left.add_run().add_picture(right_image, height=Inches(0.70))

    p_center_1 = center_cell.paragraphs[0]
    style_paragraph(p_center_1, WD_ALIGN_PARAGRAPH.LEFT)
    run = p_center_1.add_run("MSW MANAGEMENT")
    style_run(run, color=RED)
    run = p_center_1.add_run(" -Multidisciplinary, Scientific Work and Management Journal")
    style_run(run, color=BLACK, bold=False)

    p_center_2 = center_cell.add_paragraph()
    style_paragraph(p_center_2, WD_ALIGN_PARAGRAPH.LEFT)
    run = p_center_2.add_run("ISSN: 1053-7899")
    style_run(run, color=BLACK)

    p_center_3 = center_cell.add_paragraph()
    style_paragraph(p_center_3, WD_ALIGN_PARAGRAPH.LEFT)
    run = p_center_3.add_run(
        f"Vol. {volume} Issue {issue}, {paper_year}, Pages: {start_page}-{end_page}"
    )
    style_run(run, color=BLACK)

    p_right = right_cell.paragraphs[0]
    style_paragraph(p_right, WD_ALIGN_PARAGRAPH.RIGHT)
    if left_image and os.path.exists(left_image):
        p_right.add_run().add_picture(left_image, height=Inches(0.70))

    p_right_2 = right_cell.add_paragraph()
    style_paragraph(p_right_2, WD_ALIGN_PARAGRAPH.RIGHT)
    run = p_right_2.add_run("ELSEVIER")
    style_run(run, color=BLACK)


def build_header(
    section,
    start_page,
    end_page,
    left_image,
    right_image,
    *,
    volume="36",
    paper_year=PAPER_YEAR,
    issue=ISSUE,
):
    ensure_header_clearance(section)
    section.header_distance = Inches(0.15)
    section.different_first_page_header_footer = False

    for header in _iter_header_variants(section):
        header.is_linked_to_previous = False
        _populate_header_story(
            header,
            section,
            start_page,
            end_page,
            left_image,
            right_image,
            volume=volume,
            paper_year=paper_year,
            issue=issue,
        )


def _populate_header_story_ijrss(
    header,
    section,
    *,
    volume,
    issue,
    paper_year,
    paper_month,
):
    clear_story(header)

    usable_width = section.page_width - section.left_margin - section.right_margin
    table = header.add_table(rows=1, cols=1, width=usable_width)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    set_cell_border(cell, "bottom", 18, color=IJRSS_ORANGE_HEX)

    p_title = cell.paragraphs[0]
    style_paragraph(p_title, WD_ALIGN_PARAGRAPH.LEFT)
    style_run_custom(
        p_title.add_run("International Journal of Research in Social Sciences"),
        font_name="Calibri",
        color=IJRSS_BLUE,
        size=Pt(16),
        bold=True,
    )

    p_vol = cell.add_paragraph()
    style_paragraph(p_vol, WD_ALIGN_PARAGRAPH.LEFT)
    style_run_custom(
        p_vol.add_run(f"Vol. {volume} Issue {issue}, {paper_month} {paper_year}"),
        font_name="Calibri",
        color=BLACK,
        size=Pt(11),
        bold=False,
    )

    p_issn = cell.add_paragraph()
    style_paragraph(p_issn, WD_ALIGN_PARAGRAPH.LEFT)
    style_run_custom(
        p_issn.add_run(f"ISSN: {IJRSS_ISSN} Impact Factor: {IJRSS_IMPACT_FACTOR}"),
        font_name="Calibri",
        color=BLACK,
        size=Pt(11),
        bold=False,
    )

    p_home = cell.add_paragraph()
    style_paragraph(p_home, WD_ALIGN_PARAGRAPH.LEFT)
    style_run_custom(
        p_home.add_run("Journal Homepage: "),
        font_name="Calibri",
        color=BLACK,
        size=Pt(11),
        bold=False,
    )
    style_run_custom(
        p_home.add_run(IJRSS_HOMEPAGE),
        font_name="Calibri",
        color=IJRSS_BLUE,
        size=Pt(11),
        bold=False,
        underline=True,
    )
    style_run_custom(
        p_home.add_run(f", Email: {IJRSS_EMAIL}"),
        font_name="Calibri",
        color=BLACK,
        size=Pt(11),
        bold=False,
    )

    p_desc = cell.add_paragraph()
    style_paragraph(p_desc, WD_ALIGN_PARAGRAPH.LEFT)
    style_run_custom(
        p_desc.add_run(
            "Double-Blind Peer Reviewed Refereed Open Access International Journal - Included in the "
            "International Serial Directories Indexed & Listed at: Ulrich's Periodicals Directory ©, "
            "U.S.A., Open J-Gate as well as in Cabell’s Directories of Publishing Opportunities, U.S.A"
        ),
        font_name="Calibri",
        color=BLACK,
        size=Pt(11),
        bold=False,
    )


def build_header_ijrss(
    section,
    *,
    volume,
    issue,
    paper_year,
    paper_month,
):
    ensure_header_clearance(section, minimum_top_margin=Inches(1.45))
    section.header_distance = Inches(0.2)
    section.different_first_page_header_footer = False

    for header in _iter_header_variants(section):
        header.is_linked_to_previous = False
        _populate_header_story_ijrss(
            header,
            section,
            volume=volume,
            issue=issue,
            paper_year=paper_year,
            paper_month=paper_month,
        )


def _populate_header_story_ijmie(
    header,
    section,
    *,
    volume,
    issue,
    paper_year,
    paper_month,
):
    clear_story(header)

    usable_width = section.page_width - section.left_margin - section.right_margin
    table = header.add_table(rows=1, cols=1, width=usable_width)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    set_cell_border(cell, "bottom", 18, color=IJRSS_ORANGE_HEX)

    p_title = cell.paragraphs[0]
    style_paragraph(p_title, WD_ALIGN_PARAGRAPH.LEFT)
    style_run_custom(
        p_title.add_run("International Journal of Management, IT & Engineering"),
        font_name="Calibri",
        color=IJRSS_BLUE,
        size=Pt(16),
        bold=True,
    )

    p_vol = cell.add_paragraph()
    style_paragraph(p_vol, WD_ALIGN_PARAGRAPH.LEFT)
    style_run_custom(
        p_vol.add_run(f"Vol. {volume} Issue {issue}, {paper_month} {paper_year}"),
        font_name="Calibri",
        color=BLACK,
        size=Pt(11),
        bold=False,
    )

    p_issn = cell.add_paragraph()
    style_paragraph(p_issn, WD_ALIGN_PARAGRAPH.LEFT)
    style_run_custom(
        p_issn.add_run(f"ISSN: {IJMIE_ISSN} Impact Factor: {IJMIE_IMPACT_FACTOR}"),
        font_name="Calibri",
        color=BLACK,
        size=Pt(11),
        bold=False,
    )

    p_home = cell.add_paragraph()
    style_paragraph(p_home, WD_ALIGN_PARAGRAPH.LEFT)
    style_run_custom(
        p_home.add_run("Journal Homepage: "),
        font_name="Calibri",
        color=BLACK,
        size=Pt(11),
        bold=False,
    )
    style_run_custom(
        p_home.add_run(IJRSS_HOMEPAGE),
        font_name="Calibri",
        color=IJRSS_BLUE,
        size=Pt(11),
        bold=False,
        underline=True,
    )
    style_run_custom(
        p_home.add_run(f", Email: {IJRSS_EMAIL}"),
        font_name="Calibri",
        color=BLACK,
        size=Pt(11),
        bold=False,
    )

    p_desc = cell.add_paragraph()
    style_paragraph(p_desc, WD_ALIGN_PARAGRAPH.LEFT)
    style_run_custom(
        p_desc.add_run(
            "Double-Blind Peer Reviewed Refereed Open Access International Journal - Included in the "
            "International Serial Directories Indexed & Listed at: Ulrich's Periodicals Directory ©, "
            "U.S.A., Open J-Gate as well as in Cabell’s Directories of Publishing Opportunities, U.S.A"
        ),
        font_name="Calibri",
        color=BLACK,
        size=Pt(11),
        bold=False,
    )


def _populate_header_story_ijmie_inner(header, section):
    clear_story(header)

    usable_width = section.page_width - section.left_margin - section.right_margin
    table = header.add_table(rows=1, cols=1, width=usable_width)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
    set_cell_border(cell, "bottom", 6, color="000000")

    p = cell.paragraphs[0]
    style_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER)
    style_run_custom(
        p.add_run(f"ISSN: {IJMIE_ISSN} \U0001F4D6 Impact Factor: {IJMIE_IMPACT_FACTOR}"),
        font_name="Calibri",
        color=BLACK,
        size=Pt(12),
        bold=False,
    )


def build_header_ijmie(
    section,
    *,
    volume,
    issue,
    paper_year,
    paper_month,
):
    # First page: full IJMIE header. Other pages: compact ISSN/Impact header line.
    section.top_margin = Inches(0.9)
    section.header_distance = Inches(0.05)
    section.different_first_page_header_footer = True

    first = getattr(section, "first_page_header", None)
    if first is not None:
        first.is_linked_to_previous = False
        _populate_header_story_ijmie(
            first,
            section,
            volume=volume,
            issue=issue,
            paper_year=paper_year,
            paper_month=paper_month,
        )

    primary = getattr(section, "header", None)
    if primary is not None:
        primary.is_linked_to_previous = False
        _populate_header_story_ijmie_inner(primary, section)

    even = getattr(section, "even_page_header", None)
    if even is not None:
        even.is_linked_to_previous = False
        _populate_header_story_ijmie_inner(even, section)


def _populate_footer_story(footer, section):
    clear_story(footer)

    usable_width = section.page_width - section.left_margin - section.right_margin
    table = footer.add_table(rows=1, cols=3, width=usable_width)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell = table.cell(0, 0)
    center_cell = table.cell(0, 1)
    right_cell = table.cell(0, 2)

    left_w = Inches(1.25)
    right_w = Inches(1.50)
    center_w = usable_width - left_w - right_w
    if center_w < Inches(2.5):
        # Fallback to even-ish split on unusual page sizes.
        left_w = int(usable_width * 0.2)
        right_w = int(usable_width * 0.2)
        center_w = usable_width - left_w - right_w

    table.columns[0].width = left_w
    table.columns[1].width = center_w
    table.columns[2].width = right_w
    left_cell.width = left_w
    center_cell.width = center_w
    right_cell.width = right_w

    for cell in (left_cell, center_cell, right_cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        set_cell_border(cell, "top", 18)

    p_left = left_cell.paragraphs[0]
    style_paragraph(p_left, WD_ALIGN_PARAGRAPH.LEFT)
    style_run(p_left.add_run(""), color=BLACK)

    p_center = center_cell.paragraphs[0]
    style_paragraph(p_center, WD_ALIGN_PARAGRAPH.CENTER)
    style_run(p_center.add_run("https://mswmanagementj.com/"), color=BLACK)

    p_right = right_cell.paragraphs[0]
    style_paragraph(p_right, WD_ALIGN_PARAGRAPH.RIGHT)
    run = p_right.add_run()
    style_run(run, color=BLACK)
    add_page_field(run)


def build_footer(section):
    section.footer_distance = Inches(0.2)
    section.different_first_page_header_footer = False

    for footer in _iter_footer_variants(section):
        footer.is_linked_to_previous = False
        _populate_footer_story(footer, section)


def _populate_footer_story_ijrss(footer, section):
    clear_story(footer)

    usable_width = section.page_width - section.left_margin - section.right_margin
    table = footer.add_table(rows=1, cols=2, width=usable_width)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    left_w = Inches(0.85)
    right_w = usable_width - left_w
    if right_w < Inches(2.0):
        left_w = int(usable_width * 0.2)
        right_w = usable_width - left_w

    table.columns[0].width = left_w
    table.columns[1].width = right_w
    left_cell.width = left_w
    right_cell.width = right_w

    for cell in (left_cell, right_cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        set_cell_border(cell, "top", 12, color=IJRSS_GRAY_HEX)

    set_cell_border(left_cell, "right", 12, color=IJRSS_GRAY_HEX)

    p_left = left_cell.paragraphs[0]
    style_paragraph(p_left, WD_ALIGN_PARAGRAPH.LEFT)
    run = p_left.add_run()
    style_run_custom(run, color=IJRSS_BLUE, size=Pt(22), bold=True, italic=True)
    add_page_field(run)

    p_right_1 = right_cell.paragraphs[0]
    style_paragraph(p_right_1, WD_ALIGN_PARAGRAPH.RIGHT)
    style_run_custom(
        p_right_1.add_run("International Journal of Research in Social Sciences"),
        color=BLACK,
        size=Pt(12),
        bold=False,
        italic=True,
    )

    p_right_2 = right_cell.add_paragraph()
    style_paragraph(p_right_2, WD_ALIGN_PARAGRAPH.RIGHT)
    style_run_custom(
        p_right_2.add_run(IJRSS_HOMEPAGE),
        color=IJRSS_BLUE,
        size=Pt(12),
        bold=False,
        underline=True,
    )
    style_run_custom(
        p_right_2.add_run(f", Email: {IJRSS_EMAIL}"),
        color=BLACK,
        size=Pt(12),
        bold=False,
    )


def build_footer_ijrss(section):
    section.footer_distance = Inches(0.25)
    section.different_first_page_header_footer = False

    for footer in _iter_footer_variants(section):
        footer.is_linked_to_previous = False
        _populate_footer_story_ijrss(footer, section)


def _populate_footer_story_ijmie(footer, section):
    clear_story(footer)

    usable_width = section.page_width - section.left_margin - section.right_margin
    table = footer.add_table(rows=1, cols=2, width=usable_width)
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    left_w = Inches(0.85)
    right_w = usable_width - left_w
    if right_w < Inches(2.0):
        left_w = int(usable_width * 0.2)
        right_w = usable_width - left_w

    table.columns[0].width = left_w
    table.columns[1].width = right_w
    left_cell.width = left_w
    right_cell.width = right_w

    for cell in (left_cell, right_cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        set_cell_border(cell, "top", 12, color=IJRSS_GRAY_HEX)

    set_cell_border(left_cell, "right", 12, color=IJRSS_GRAY_HEX)

    p_left = left_cell.paragraphs[0]
    style_paragraph(p_left, WD_ALIGN_PARAGRAPH.LEFT)
    run = p_left.add_run()
    style_run_custom(run, color=IJRSS_BLUE, size=Pt(22), bold=True, italic=True)
    add_page_field(run)

    p_right_1 = right_cell.paragraphs[0]
    style_paragraph(p_right_1, WD_ALIGN_PARAGRAPH.RIGHT)
    style_run_custom(
        p_right_1.add_run("International journal of Management, IT and Engineering"),
        color=BLACK,
        size=Pt(12),
        bold=False,
        italic=True,
    )

    p_right_2 = right_cell.add_paragraph()
    style_paragraph(p_right_2, WD_ALIGN_PARAGRAPH.RIGHT)
    style_run_custom(
        p_right_2.add_run(IJRSS_HOMEPAGE),
        color=IJRSS_BLUE,
        size=Pt(12),
        bold=False,
        underline=True,
    )
    style_run_custom(
        p_right_2.add_run(f", Email: {IJRSS_EMAIL}"),
        color=BLACK,
        size=Pt(12),
        bold=False,
    )


def build_footer_ijmie(section):
    section.footer_distance = Inches(0.25)
    section.different_first_page_header_footer = True

    for footer in _iter_footer_variants(section):
        footer.is_linked_to_previous = False
        _populate_footer_story_ijmie(footer, section)


def apply_font(doc, *, font_name="Times New Roman", font_size=Pt(9)):
    clean_blank_paragraphs(doc)
    def iter_block_items(parent):
        if isinstance(parent, DocxDocument):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise TypeError(f"Unsupported parent type: {type(parent)!r}")

        for child in parent_elm.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, parent)
            elif child.tag.endswith("}tbl"):
                yield Table(child, parent)

    def iter_paragraphs(parent):
        for item in iter_block_items(parent):
            if isinstance(item, Paragraph):
                yield item
            else:
                for row in item.rows:
                    for cell in row.cells:
                        yield from iter_paragraphs(cell)

    paragraphs = list(iter_paragraphs(doc))

    for paragraph in paragraphs:
        style_paragraph(paragraph, paragraph.alignment)
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = font_size


def format_docx_files(
    input_paths,
    output_dir,
    *,
    config=FormatConfig(),
    reference_dir=None,
):
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    template = (getattr(config, "template", None) or "msw").strip().lower()
    if template == "msw":
        left_image, right_image = ensure_reference_images(str(output_dir), reference_dir=reference_dir)
    else:
        left_image, right_image = "", ""

    current_page = int(config.start_page)
    output_paths = []
    for input_path in input_paths:
        input_path = Path(input_path)
        if input_path.suffix.lower() != ".docx":
            raise ValueError(f"Only .docx is supported (got: {input_path.name})")

        page_count = get_page_count(str(input_path))
        end_page = current_page + page_count - 1

        doc = Document(str(input_path))
        doc.settings.odd_and_even_pages_header_footer = False
        if template in ("ijrss", "ijmie"):
            apply_font(doc, font_size=Pt(12))
        else:
            apply_font(doc, font_size=Pt(9))
        for section in doc.sections:
            if template == "ijrss":
                build_header_ijrss(
                    section,
                    volume=str(config.volume),
                    issue=str(config.issue),
                    paper_year=str(config.paper_year),
                    paper_month=str(getattr(config, "paper_month", "March") or "March"),
                )
                build_footer_ijrss(section)
            elif template == "ijmie":
                build_header_ijmie(
                    section,
                    volume=str(config.volume),
                    issue=str(config.issue),
                    paper_year=str(config.paper_year),
                    paper_month=str(getattr(config, "paper_month", "March") or "March"),
                )
                build_footer_ijmie(section)
            elif template == "msw":
                build_header(
                    section,
                    current_page,
                    end_page,
                    left_image,
                    right_image,
                    volume=str(config.volume),
                    paper_year=str(config.paper_year),
                    issue=str(config.issue),
                )
                build_footer(section)
            else:
                raise ValueError(f"Unknown template: {template!r} (expected 'msw', 'ijrss', or 'ijmie')")
        set_starting_page_number(doc, current_page)

        out_path = output_dir / input_path.name
        doc.save(str(out_path))
        output_paths.append(str(out_path))

        current_page = end_page + 1

    return output_paths


def format_papers(directory, start_page=START_PAGE):
    left_image, right_image = ensure_reference_images(directory, reference_dir=directory)

    files = sorted(
        [name for name in os.listdir(directory) if name.startswith("Paper ") and name.endswith(".docx")],
        key=lambda name: int(re.search(r"Paper (\d+)", name).group(1)),
    )

    current_page = start_page
    for idx, filename in enumerate(files):
        path = os.path.join(directory, filename)
        page_count = get_page_count(path)

        end_page = current_page + page_count - 1

        doc = Document(path)
        doc.settings.odd_and_even_pages_header_footer = False
        apply_font(doc)
        for section in doc.sections:
            build_header(
                section,
                current_page,
                end_page,
                left_image,
                right_image,
                volume="36",
                paper_year=PAPER_YEAR,
                issue=ISSUE,
            )
            build_footer(section)
        set_starting_page_number(doc, current_page)
        doc.save(path)

        current_page = end_page + 1


if __name__ == "__main__":
    format_papers("/Users/apple/Documents/wordd", start_page=START_PAGE)
